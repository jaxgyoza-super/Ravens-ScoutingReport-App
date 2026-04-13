import io
from concurrent.futures import ThreadPoolExecutor as _TPE
from datetime import date

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

import analyzer as az
import chart_generator as cg
import data_loader as dl
import auto_analyzer as aa

# ── OFF FORM 横並びペアリング定義 ──────────────────────────────
_OFF_FORM_PAIRS = [
    ('ACE',               'SUPER'),
    ('WAC',               'WAT'),
    ('SPREAD',            'DISH'),
    ('SLOT（A/B/VEER合算）', 'EMPTY'),
    ('BUNCH合算',          None),
    ('PRO合算（I/B/A）',   None),
    ('PRO FAR',           'PRO NEAR'),
    ('TWIN合算（A/B/I）',  None),
    ('TWIN FAR',          'TWIN NEAR'),
    ('UMB系（合算）',      None),
]


# ── スタイルヘルパー ────────────────────────────────────────
def _set_cell_font(cell, size_pt: float, bold: bool = False):
    """セル内の全テキストにフォントサイズ・游ゴシック Medium・太字を適用"""
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(size_pt)
            _set_jpfont(run, '游ゴシック Medium')
            if bold:
                run.bold = True
        # run が無い場合（直接 cell.text で入れた場合）も対応
        if not para.runs and para.text:
            run = para.add_run(para.text)
            para.clear()
            run = para.add_run(para.text)
            run.font.size = Pt(size_pt)
            _set_jpfont(run, '游ゴシック Medium')


_TBL_REPLACE = [
    ('関西学院', '関学'),
    ('立命館',   '立命'),
    ('SIGN(D)',  'SIGN'),
    ('SIGN(d)',  'SIGN'),
]

_COMP_LEGACY = {'GREEN': 'FS Sky', 'SILVER': 'FS Buzz', 'GOLD': 'R Buzz', 'YELLOW': 'R Sky'}

def _norm(text: str) -> str:
    """テーブルセル内の表記を統一する"""
    for old, new in _TBL_REPLACE:
        text = text.replace(old, new)
    # SIGN値「N」（ノーブリッツ）はセル内容がそのまま'N'の場合のみ'read'に変換
    if text == 'N':
        return 'read'
    # COMPONENT旧表記を統一（セル値が正確に一致する場合のみ）
    if text in _COMP_LEGACY:
        return _COMP_LEGACY[text]
    return text


def _keep_table_together(table):
    """テーブル全体を同一ページに収める（行分割禁止 + 各行のヘッダー段落 keep_with_next）"""
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement('w:cantSplit')
        cant.set(qn('w:val'), '1')
        trPr.append(cant)


def _keep_block_together(para, table):
    """見出し段落 + テーブルを同一ページに収める"""
    para.paragraph_format.keep_with_next = True
    _keep_table_together(table)
    # テーブルの最初の行ヘッダー段落も keep_with_next で見出しと繋げる
    if table.rows:
        first_row = table.rows[0]
        for cell in first_row.cells:
            for p in cell.paragraphs:
                if p.text.strip():
                    p.paragraph_format.keep_with_next = True
                    break


def _set_col_widths_mm(table, widths_mm):
    """
    各列幅を mm 単位で確実に設定する。
    table.columns[x].width は Word に無視されることがあるため、
    各セルの w:tcW を直接書き込む。
    """
    twips_per_mm = 1440 / 25.4
    total_twips  = int(sum(widths_mm) * twips_per_mm)

    # テーブル全体幅も明示
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(total_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # 各セルに幅を書き込む
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i >= len(widths_mm):
                break
            w_twips = int(widths_mm[i] * twips_per_mm)
            tcPr = cell._tc.get_or_add_tcPr()
            for old in tcPr.findall(qn('w:tcW')):
                tcPr.remove(old)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'),    str(w_twips))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)


def _set_picture_float_square(doc, width_emu=None):
    """直前に追加された画像段落を四角形の文字列折り返し（浮動）に変換する"""
    para = doc.paragraphs[-1]
    drawing_elem = None
    for run in para.runs:
        d = run._r.find(qn('w:drawing'))
        if d is not None:
            drawing_elem = d
            break
    if drawing_elem is None:
        return

    inline = drawing_elem.find(qn('wp:inline'))
    if inline is None:
        return

    # 現在のサイズを取得
    extent = inline.find(qn('wp:extent'))
    cx = int(extent.get('cx'))
    cy = int(extent.get('cy'))
    if width_emu and width_emu != cx:
        cy = int(cy * width_emu / cx)
        cx = width_emu

    # wp:anchor を構築
    anchor = OxmlElement('wp:anchor')
    for k, v in [('distT', '0'), ('distB', '0'), ('distL', '0'), ('distR', '0'),
                 ('simplePos', '0'), ('relativeHeight', '251658240'),
                 ('behindDoc', '0'), ('locked', '0'),
                 ('layoutInCell', '1'), ('allowOverlap', '0')]:
        anchor.set(k, v)

    sp = OxmlElement('wp:simplePos')
    sp.set('x', '0'); sp.set('y', '0')
    anchor.append(sp)

    posH = OxmlElement('wp:positionH')
    posH.set('relativeFrom', 'margin')
    align_h = OxmlElement('wp:align')
    align_h.text = 'left'
    posH.append(align_h)
    anchor.append(posH)

    posV = OxmlElement('wp:positionV')
    posV.set('relativeFrom', 'paragraph')
    ov = OxmlElement('wp:posOffset')
    ov.text = '0'
    posV.append(ov)
    anchor.append(posV)

    new_extent = OxmlElement('wp:extent')
    new_extent.set('cx', str(cx)); new_extent.set('cy', str(cy))
    anchor.append(new_extent)

    ee = inline.find(qn('wp:effectExtent'))
    if ee is None:
        ee = OxmlElement('wp:effectExtent')
        for k in ('l', 't', 'r', 'b'):
            ee.set(k, '0')
    anchor.append(ee)

    wrap = OxmlElement('wp:wrapSquare')
    wrap.set('wrapText', 'bothSides')
    anchor.append(wrap)

    for tag in ('wp:docPr', 'wp:cNvGraphicFramePr'):
        el = inline.find(qn(tag))
        if el is not None:
            anchor.append(el)

    # a:graphic（名前空間付き検索）
    _a_ns = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    graphic = inline.find(f'{{{_a_ns}}}graphic')
    if graphic is None:
        graphic = inline.find(qn('a:graphic'))
    if graphic is not None:
        anchor.append(graphic)

    drawing_elem.remove(inline)
    drawing_elem.append(anchor)


def _set_jpfont(run, font_name: str):
    """ランに日本語フォントを設定（ASCII + eastAsia 両方）"""
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def _tight_spacing(p):
    """段落の行間を1.0・段落前後スペースを0ptに設定"""
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(0)
    fmt.space_after  = Pt(0)


def _apply_table_font(table, default_pt: float = 12, biko_pt: float = 10.0):
    """
    テーブル全体に default_pt・游ゴシック Medium を適用。
    備考列（3列目）のデータ行：
      - 大学名＋プレー番号行（'#数字' を含む）→ biko_pt（10pt）
      - 説明文行（それ以外）                   → default_pt（12pt）
    ヘッダー行（row 0）は常に default_pt + 太字。
    """
    import re as _re_font
    has_biko = len(table.columns) == 3
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            is_header = (r_idx == 0)
            is_biko   = has_biko and (c_idx == 2) and not is_header
            for para in cell.paragraphs:
                # 備考列：行内容で字サイズを切り替える
                if is_biko:
                    size = biko_pt if _re_font.search(r'#\d+', para.text) else default_pt
                else:
                    size = default_pt
                for run in para.runs:
                    run.font.size = Pt(size)
                    _set_jpfont(run, '游ゴシック Medium')
                    if is_header:
                        run.bold = True


def _set_row_bold(row):
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True


def _shade_cell(cell, hex_color='D9E1F2'):
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """セルに枠線を設定"""
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '12')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), 'auto')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)


def _format_table(table):
    """テーブルを見やすくフォーマット"""
    from docx.shared import Inches
    from docx.oxml import OxmlElement
    # 列幅設定
    if len(table.columns) == 3:
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(2.5)
    # 全セルに枠線と余白
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top={}, left={}, bottom={}, right={})
            # セル内の余白
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_name in ['top', 'left', 'bottom', 'right']:
                m = OxmlElement(f'w:{margin_name}')
                m.set(qn('w:w'), '100')
                m.set(qn('w:type'), 'dxa')
                tcMar.append(m)
            tcPr.append(tcMar)
    # フォントサイズ適用（本文12pt・備考列10.5pt）
    _apply_table_font(table)


def _add_coverage_with_comp3(doc, main_tbl, comp3_tbl, title=None, auto_bullets=None):
    """COVERAGE表にCover3内訳行を埋め込んだ1つのテーブルを生成"""
    if main_tbl is None or len(main_tbl) == 0:
        if title:
            doc.add_paragraph(f'{title}：該当データなし').runs[0].italic = True
        return

    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(18)

    # 自動考察：タイトルの下・テーブルの上に挿入
    if auto_bullets:
        _add_auto_section(doc, auto_bullets)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # ヘッダー行
    hdr = table.rows[0].cells
    for i, name in enumerate(['COVERAGE', '割合（実数）', '備考']):
        hdr[i].text = name
        _shade_cell(hdr[i], 'D9E1F2')
    _set_row_bold(table.rows[0])

    # comp3をdict化（COMPONENTがキー）
    comp3_rows = []
    if comp3_tbl is not None and len(comp3_tbl) > 0:
        for _, r in comp3_tbl.iterrows():
            comp3_rows.append(r)

    for _, row in main_tbl.iterrows():
        cov_val = str(row.iloc[0]) if not pd.isna(row.iloc[0]) else ''
        cells = table.add_row().cells
        cells[0].text = _norm(cov_val)
        cells[1].text = _norm(str(row.iloc[1])) if not pd.isna(row.iloc[1]) else ''
        cells[2].text = _norm(str(row.iloc[2])) if not pd.isna(row.iloc[2]) else ''

        # COVERAGE = "3" の直後にCOMPONENT内訳行を挿入
        if cov_val == '3' and comp3_rows:
            for cr in comp3_rows:
                sub_cells = table.add_row().cells
                comp_name = str(cr.iloc[0]) if not pd.isna(cr.iloc[0]) else ''
                sub_cells[0].text = f'　{_norm(comp_name)}'
                sub_cells[1].text = _norm(str(cr.iloc[1])) if not pd.isna(cr.iloc[1]) else ''
                sub_cells[2].text = _norm(str(cr.iloc[2])) if not pd.isna(cr.iloc[2]) else ''
                for cell in sub_cells:
                    _shade_cell(cell, 'E2EFDA')

    _format_table(table)
    doc.add_paragraph()


def _clear_cell_borders(cell):
    """外枠テーブルのセル枠線を非表示にする"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _add_coverage_to_cell(doc, cell, main_tbl, comp3_tbl,
                          grp_name=None, three_var_dict=None):
    """Coverage テーブルを生成してcell内に配置（ネストテーブル）"""
    if main_tbl is None or len(main_tbl) == 0:
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for i, name in enumerate(['COVERAGE', '割合（実数）', '備考']):
        hdr[i].text = name
        _shade_cell(hdr[i], 'D9E1F2')
    _set_row_bold(table.rows[0])

    comp3_rows = [r for _, r in comp3_tbl.iterrows()] if (
        comp3_tbl is not None and len(comp3_tbl) > 0) else []

    for _, row in main_tbl.iterrows():
        cov_val = str(row.iloc[0]) if not pd.isna(row.iloc[0]) else ''
        cells = table.add_row().cells
        cells[0].text = _norm(cov_val)
        cells[1].text = _norm(str(row.iloc[1])) if not pd.isna(row.iloc[1]) else ''
        biko_text = _norm(str(row.iloc[2])) if not pd.isna(row.iloc[2]) else ''

        # 3変数複合考察を備考に追記
        if three_var_dict is not None and grp_name and cov_val:
            extra = three_var_dict.get((str(grp_name), str(cov_val)), [])
            if extra:
                extra_text = ' / '.join(extra)
                biko_text = f'{biko_text}　{extra_text}' if biko_text else extra_text

        cells[2].text = biko_text

        if cov_val == '3' and comp3_rows:
            for cr in comp3_rows:
                sub = table.add_row().cells
                sub[0].text = f'　{_norm(str(cr.iloc[0]))}'
                sub[1].text = _norm(str(cr.iloc[1])) if not pd.isna(cr.iloc[1]) else ''
                sub[2].text = _norm(str(cr.iloc[2])) if not pd.isna(cr.iloc[2]) else ''
                for c in sub:
                    _shade_cell(c, 'E2EFDA')

    # 列幅を XML レベルで確実に設定（25mm / 30mm / 21mm）
    _set_col_widths_mm(table, [25, 30, 21])
    for row in table.rows:
        for c in row.cells:
            _set_cell_border(c, top={}, left={}, bottom={}, right={})
    _apply_table_font(table)

    # ネスト: doc.bodyから切り離してcellに配置
    tbl_xml = table._tbl
    doc.element.body.remove(tbl_xml)
    cell._tc.append(tbl_xml)
    # OOXML仕様: w:tc は必ず w:p で終わらなければならない（ないと破損扱い）
    cell._tc.append(OxmlElement('w:p'))


def _add_sign_d_to_cell(doc, cell, sign_tbl):
    """SIGN(D)テーブルをcell内に配置（ネストテーブル）"""
    if sign_tbl is None or len(sign_tbl) == 0:
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for i, name in enumerate(['SIGN', '割合（実数）', '備考']):
        hdr[i].text = name
        _shade_cell(hdr[i], 'D9E1F2')
    _set_row_bold(table.rows[0])

    for _, row in sign_tbl.iterrows():
        cells = table.add_row().cells
        cells[0].text = _norm(str(row.iloc[0])) if not pd.isna(row.iloc[0]) else ''
        cells[1].text = _norm(str(row.iloc[1])) if not pd.isna(row.iloc[1]) else ''
        cells[2].text = _norm(str(row.iloc[2])) if not pd.isna(row.iloc[2]) else ''

    _set_col_widths_mm(table, [25, 30, 21])
    for row in table.rows:
        for c in row.cells:
            _set_cell_border(c, top={}, left={}, bottom={}, right={})
    _apply_table_font(table)

    tbl_xml = table._tbl
    doc.element.body.remove(tbl_xml)
    cell._tc.append(tbl_xml)
    cell._tc.append(OxmlElement('w:p'))


def _set_cell_multiline(cell, parts):
    """セルに複数行テキストを・区切りで設定する"""
    if not parts:
        cell.paragraphs[0].clear()
        return
    first = cell.paragraphs[0]
    first.clear()
    r = first.add_run(f'・{parts[0]}')
    r.font.size = Pt(10)
    for part in parts[1:]:
        p = cell.add_paragraph()
        r = p.add_run(f'・{part}')
        r.font.size = Pt(10)


def _add_off_form_sign_d_paired(doc, grp_results):
    """OFF FORM × SIGN(D) を定義済みペアで2列横並びに表示"""
    CELL_W = Cm(8.0)

    grp_dict = {g[0]: g for g in grp_results}
    shown = set()

    def _emit_pair(left_grp, right_grp):
        outer = doc.add_table(rows=1, cols=2)
        for cell in outer.rows[0].cells:
            _clear_cell_borders(cell)
        for col in outer.columns:
            col.width = CELL_W

        for col_idx, grp in enumerate([left_grp, right_grp]):
            cell = outer.cell(0, col_idx)
            if grp is None:
                continue
            grp_name, n, sign_tbl = grp
            p = cell.paragraphs[0]
            run = p.add_run(f'{grp_name}（{n}プレー）')
            run.bold = True
            _add_sign_d_to_cell(doc, cell, sign_tbl)

        _keep_table_together(outer)
        doc.add_paragraph()

    for left_name, right_name in _OFF_FORM_PAIRS:
        left  = grp_dict.get(left_name)
        right = grp_dict.get(right_name) if right_name else None
        if left is None and right is None:
            continue
        if left_name:
            shown.add(left_name)
        if right_name:
            shown.add(right_name)
        _emit_pair(left, right)

    others = [g for g in grp_results if g[0] not in shown]
    for i in range(0, len(others), 2):
        left  = others[i]
        right = others[i + 1] if i + 1 < len(others) else None
        _emit_pair(left, right)


def _add_cov_to_cell(doc, cell, grp_name, cov_tbl, comp3_tbl, three_var_dict):
    """COVERAGE テーブルをセル内にネスト配置（その他隊形の横並び用）"""
    if cov_tbl is None or len(cov_tbl) == 0:
        return

    comp3_rows = [row for _, row in comp3_tbl.iterrows()] if (
        comp3_tbl is not None and len(comp3_tbl) > 0) else []

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for i, name in enumerate(['COVERAGE', '割合（実数）', '備考']):
        hdr[i].text = name
        _shade_cell(hdr[i], 'D9E1F2')
    _set_row_bold(table.rows[0])

    for _, row in cov_tbl.iterrows():
        cov_val = str(row.iloc[0]) if not pd.isna(row.iloc[0]) else ''
        cells = table.add_row().cells
        cells[0].text = _norm(cov_val)
        cells[1].text = _norm(str(row.iloc[1])) if not pd.isna(row.iloc[1]) else ''

        biko_parts = []
        base_biko = _norm(str(row.iloc[2])) if not pd.isna(row.iloc[2]) and str(row.iloc[2]).strip() not in ('', 'nan') else ''
        if base_biko:
            biko_parts.append(base_biko)
        if three_var_dict and grp_name and cov_val:
            biko_parts.extend(three_var_dict.get((str(grp_name), str(cov_val)), []))
        _set_cell_multiline(cells[2], biko_parts)

        if cov_val == '3' and comp3_rows:
            for cr in comp3_rows:
                sub = table.add_row().cells
                sub[0].text = f'　{_norm(str(cr.iloc[0]))}'
                sub[1].text = _norm(str(cr.iloc[1])) if not pd.isna(cr.iloc[1]) else ''
                sub_biko = _norm(str(cr.iloc[2])) if not pd.isna(cr.iloc[2]) and str(cr.iloc[2]).strip() not in ('', 'nan') else ''
                sub[2].text = sub_biko
                for c in sub:
                    _shade_cell(c, 'E2EFDA')

    # セル幅: SIGN(D) 横並び表に合わせて [25, 30, 21] = 76mm
    _set_col_widths_mm(table, [25, 30, 21])
    for row in table.rows:
        for c in row.cells:
            _set_cell_border(c, top={}, left={}, bottom={}, right={})
    _apply_table_font(table)

    tbl_xml = table._tbl
    doc.element.body.remove(tbl_xml)
    cell._tc.append(tbl_xml)
    cell._tc.append(OxmlElement('w:p'))


def _add_off_form_paired(doc, grp_results, three_var_dict=None):
    """OFF FORM × COVERAGE を表示。
    - 定義済み隊形：縦並び・最大幅
    - その他隊形：横2列並び（_add_off_form_sign_d_paired と同方式）
    """
    # 定義済み隊形名セット（_OFF_FORM_PAIRS から構築）
    _defined = {name for pair in _OFF_FORM_PAIRS for name in pair if name}

    defined_grps = [(n, c, t, p) for n, c, t, p in grp_results if n in _defined]
    other_grps   = [(n, c, t, p) for n, c, t, p in grp_results if n not in _defined]

    # ── 定義済み：縦並び・最大幅 ──────────────────────────────
    for grp_name, n, cov_tbl, comp3_tbl in defined_grps:
        p = doc.add_paragraph()
        r = p.add_run(f'{grp_name}（{n}プレー）')
        r.bold = True
        r.font.size = Pt(12)

        if cov_tbl is None or len(cov_tbl) == 0:
            p2 = doc.add_paragraph('  該当データなし')
            p2.runs[0].italic = True
            doc.add_paragraph()
            continue

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        for i, name in enumerate(['COVERAGE', '割合（実数）', '備考']):
            hdr[i].text = name
            _shade_cell(hdr[i], 'D9E1F2')
        _set_row_bold(table.rows[0])

        comp3_rows = [row for _, row in comp3_tbl.iterrows()] if (
            comp3_tbl is not None and len(comp3_tbl) > 0) else []

        for _, row in cov_tbl.iterrows():
            cov_val = str(row.iloc[0]) if not pd.isna(row.iloc[0]) else ''
            cells = table.add_row().cells
            cells[0].text = _norm(cov_val)
            cells[1].text = _norm(str(row.iloc[1])) if not pd.isna(row.iloc[1]) else ''

            biko_parts = []
            base_biko = _norm(str(row.iloc[2])) if not pd.isna(row.iloc[2]) and str(row.iloc[2]).strip() not in ('', 'nan') else ''
            if base_biko:
                biko_parts.append(base_biko)
            if three_var_dict and grp_name and cov_val:
                biko_parts.extend(three_var_dict.get((str(grp_name), str(cov_val)), []))
            _set_cell_multiline(cells[2], biko_parts)

            if cov_val == '3' and comp3_rows:
                for cr in comp3_rows:
                    sub = table.add_row().cells
                    sub[0].text = f'　{_norm(str(cr.iloc[0]))}'
                    sub[1].text = _norm(str(cr.iloc[1])) if not pd.isna(cr.iloc[1]) else ''
                    sub_biko = _norm(str(cr.iloc[2])) if not pd.isna(cr.iloc[2]) and str(cr.iloc[2]).strip() not in ('', 'nan') else ''
                    sub[2].text = sub_biko
                    for c in sub:
                        _shade_cell(c, 'E2EFDA')

        # 最大幅: [28, 33, 99] = 160mm
        _set_col_widths_mm(table, [28, 33, 99])
        for row in table.rows:
            for c in row.cells:
                _set_cell_border(c, top={}, left={}, bottom={}, right={})
        _apply_table_font(table)
        _keep_block_together(p, table)
        doc.add_paragraph()

    # ── その他隊形：横2列並び ────────────────────────────────
    CELL_W = Cm(8.0)
    for i in range(0, len(other_grps), 2):
        left_grp  = other_grps[i]
        right_grp = other_grps[i + 1] if i + 1 < len(other_grps) else None

        outer = doc.add_table(rows=1, cols=2)
        for cell in outer.rows[0].cells:
            _clear_cell_borders(cell)
        for col in outer.columns:
            col.width = CELL_W

        for col_idx, grp in enumerate([left_grp, right_grp]):
            if grp is None:
                continue
            grp_name, n, cov_tbl, comp3_tbl = grp
            cell = outer.cell(0, col_idx)
            p = cell.paragraphs[0]
            run = p.add_run(f'{grp_name}（{n}プレー）')
            run.bold = True
            _add_cov_to_cell(doc, cell, grp_name, cov_tbl, comp3_tbl, three_var_dict)

        _keep_table_together(outer)
        doc.add_paragraph()


def _add_table(doc, df, title=None, header_color='D9E1F2', indent_level=0, auto_bullets=None):
    """DataFrameをWordテーブルとして追記する"""
    if df is None or len(df) == 0:
        if title:
            p = doc.add_paragraph(f'  {'　' * indent_level}{title}：該当データなし')
            p.runs[0].italic = True
        return

    if title:
        p = doc.add_paragraph()
        indent_pt = indent_level * 360  # 1レベル = 0.5インチ相当
        p.paragraph_format.left_indent = indent_pt
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(18)

    # 自動考察：タイトルの下・テーブルの上に挿入
    if auto_bullets:
        _add_auto_section(doc, auto_bullets)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # ヘッダー行
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        _shade_cell(hdr_cells[i], header_color)
    _set_row_bold(table.rows[0])

    # データ行
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = '' if pd.isna(val) else _norm(str(val))

    _format_table(table)
    doc.add_paragraph()  # 余白


# ── 見出しヘルパー ──────────────────────────────────────────
_BM_ID = [0]  # ブックマークID採番用

def _add_bookmark(p, name):
    """段落の先頭にブックマークを挿入"""
    bm_id = str(_BM_ID[0])
    _BM_ID[0] += 1
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), bm_id)
    bm_start.set(qn('w:name'), name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), bm_id)
    p._p.insert(0, bm_start)
    p._p.append(bm_end)


def _add_h1(doc, text, bookmark=None):
    """大見出し（Normal Situation / 3rd / Red Zone / 2MIN など）: 28pt 赤 太字"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)
    if bookmark:
        _add_bookmark(p, bookmark)


_TOC_ENTRIES = [
    ('Normal Situation', 'bm_normal'),
    ('3rd Situation',    'bm_3rd'),
    ('Red Zone',         'bm_red'),
    ('2MIN',             'bm_2min'),
    ('選手の特徴',       'bm_players'),
]


def _add_toc(doc):
    """表紙用目次（セクション名 + ページ番号プレースホルダー）を挿入。
    返値: {bookmark_name: run} — 後でページ番号を書き込む用"""
    p_title = doc.add_paragraph()
    r = p_title.add_run('目次')
    r.bold = True
    r.font.size = Pt(16)
    _set_jpfont(r, '游ゴシック Medium')
    p_title.paragraph_format.space_before = Pt(14)
    p_title.paragraph_format.space_after  = Pt(6)

    toc_runs = {}
    for label, bm_name in _TOC_ENTRIES:
        p = doc.add_paragraph()
        r_prefix = p.add_run('・')
        r_prefix.font.size = Pt(13)
        _set_jpfont(r_prefix, '游ゴシック Medium')
        r_label = p.add_run(f'{label}　')
        r_label.font.size = Pt(13)
        _set_jpfont(r_label, '游ゴシック Medium')
        r_page = p.add_run('?ページ')
        r_page.font.size = Pt(13)
        _set_jpfont(r_page, '游ゴシック Medium')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        toc_runs[bm_name] = r_page
    return toc_runs


def _add_page_number_footer(doc):
    """全セクションのフッター右下に自動ページ番号を追加"""
    for section in doc.sections:
        section.footer_distance = Pt(14)
        footer = section.footer
        footer.is_linked_to_previous = False
        # 既存段落を使う（footerには常に1段落ある）
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        _set_jpfont(run, '游ゴシック Medium')
        run.font.size = Pt(10)
        # PAGE フィールド
        fc_begin = OxmlElement('w:fldChar')
        fc_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fc_begin)
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run._r.append(instr)
        fc_sep = OxmlElement('w:fldChar')
        fc_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fc_sep)
        t = OxmlElement('w:t')
        t.text = '1'
        run._r.append(t)
        fc_end = OxmlElement('w:fldChar')
        fc_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fc_end)


def _estimate_page_nums(doc, bookmark_names):
    """明示的改ページ（page_break_before / w:br type=page）を数え
    各ブックマーク到達時のページ番号を返す"""
    bm_set = set(bookmark_names)
    page = 1
    found = {}

    for para in doc.paragraphs:
        # この段落が改ページ前始まりなら先にページを進める
        if para.paragraph_format.page_break_before:
            page += 1
        else:
            for br in para._p.iter(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    page += 1

        # ブックマーク検出
        for bm in para._p.iter(qn('w:bookmarkStart')):
            name = bm.get(qn('w:name'))
            if name in bm_set and name not in found:
                found[name] = page

    return found


def _add_h2(doc, text):
    """小見出し（Run Defense / Pass Defense など）: 26pt 黒 太字"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)


def _add_h3(doc, text):
    """ゾーン見出し（①SHORT / ②MIDDLE① など）: 18pt 黒 太字"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)


# ── 特徴書き込みスペース ─────────────────────────────────────
_CIRCLE = '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮'

def _add_tokuchou(doc, title, df_table, named_count, total_count=None,
                  name_placeholder=None, consider_fn=None, show_fig_placeholder=True,
                  extra_names=None):
    """
    後から特徴を書き込むスペースを生成する。
    - named_count:       番号＋名前を付けるスロット数
    - total_count:       合計スロット数（None の場合は names の実数で自動決定）
    - name_placeholder:  文字列を指定すると全スロットにその文字列を使用（実データを使わない）
    - consider_fn:       関数 (value_str) -> [bullet_str, ...] を渡すと自動考察を挿入
    - extra_names:       末尾に追加する名前リスト（重複除外・total_count を自動延長）
    2行の書き込みスペース（空行）を各項目の後に挿入する。
    """
    # セクションタイトル
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    # 名前リストを構築
    if name_placeholder is not None:
        names = [name_placeholder] * named_count
    else:
        names = []
        if df_table is not None and len(df_table) > 0:
            for i, (_, row) in enumerate(df_table.iterrows()):
                if i >= named_count:
                    break
                val = str(row.iloc[0]) if not pd.isna(row.iloc[0]) else ''
                names.append(_norm(val))

    # extra_names を重複なく追加
    if extra_names:
        _seen = set(names)
        for en in extra_names:
            if en not in _seen:
                names.append(en)
                _seen.add(en)

    # 実際のスロット数（total_count 未指定なら names の実数、指定あれば最大を取る）
    actual_total = len(names) if total_count is None else max(total_count, len(names))

    for i in range(actual_total):
        circle = _CIRCLE[i] if i < len(_CIRCLE) else f'({i+1})'
        label  = f'{circle}{names[i]}' if i < len(names) else circle

        p = doc.add_paragraph()
        _tight_spacing(p)
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(14)

        # 自動考察（consider_fn が指定されており、名前が存在する場合のみ）
        if consider_fn is not None and i < len(names) and names[i]:
            try:
                auto_bullets = consider_fn(names[i])
            except Exception:
                auto_bullets = []
            if auto_bullets:
                for b in auto_bullets:
                    p2 = doc.add_paragraph()
                    _tight_spacing(p2)
                    run2 = p2.add_run(f'・{b}')
                    run2.font.size = Pt(12)
                    _set_jpfont(run2, '游ゴシック Medium')
                    run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            else:
                p2 = doc.add_paragraph()
                _tight_spacing(p2)
                run2 = p2.add_run('（自動考察なし）')
                run2.font.size = Pt(9)
                run2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        # 図挿入プレースホルダー（1行）+ 下1行スペース
        if show_fig_placeholder:
            p_fig = doc.add_paragraph()
            r_fig = p_fig.add_run('図を挿入👇')
            r_fig.font.size = Pt(11)
            r_fig.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            doc.add_paragraph()


_HEATMAP_CAPTION = {
    'normal': (
        'Normal（DN=1・2、2MIN除外、Redzone除外）において、\n'
        'フィールド位置（YARD LINE）とDown＆Distanceごとの各マスで、最も多く使われた守備指標（COVERAGE / ラッシュ人数 / SIGN / BLITZ）を表示している\n'
        '各セルの数値は「最も多く出たプレーの回数 / そのマスの総プレー数」を示す\n'
        'また、出現割合が高いほど赤く、低いほど白く色分けされる。'
    ),
    'third': (
        '3rd（DN=3・4、2MIN除外、Redzone除外）において、\n'
        'フィールド位置（YARD LINE）とDown＆Distanceごとの各マスで、最も多く使われた守備指標（COVERAGE / ラッシュ人数 / SIGN / BLITZ）を表示している\n'
        '各セルの数値は「最も多く出たプレーの回数 / そのマスの総プレー数」を示す\n'
        'また、出現割合が高いほど赤く、低いほど白く色分けされる。'
    ),
}


def _add_heatmap_caption(doc, situation: str):
    """ヒートマップ画像の直下に説明文を追加する（Normal/3rd 各1回のみ）"""
    text = _HEATMAP_CAPTION.get(situation, '')
    if not text:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    _set_jpfont(run, 'Meiryo UI')
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _add_auto_section(doc, bullets):
    """自動考察テキストリストを箇条書きでドキュメントに追記"""
    if not bullets:
        return
    for b in bullets:
        p = doc.add_paragraph()
        _tight_spacing(p)
        run = p.add_run(f'・{b}')
        run.font.size = Pt(12)
        _set_jpfont(run, '游ゴシック Medium')
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    doc.add_paragraph()


def _3rd_notable_names(df_3_sub, df_n, col, already_named, diff_thr=0.20, min_n=2):
    """
    3rd ゾーンで Normal と比べて著しく多い（rate差 diff_thr 以上）、
    または 3rd にしか出ない値を返す（already_named を除く）。
    """
    if len(df_3_sub) == 0 or col not in df_3_sub.columns:
        return []
    total_3 = len(df_3_sub)
    total_n = len(df_n)
    vc_3 = df_3_sub[col].value_counts()
    vc_n = df_n[col].value_counts() if total_n > 0 else pd.Series(dtype=int)
    seen = set(already_named)
    extras = []
    for val, cnt_3 in vc_3.items():
        if not val or str(val) in ('', 'nan'):
            continue
        norm_val = _norm(str(val))
        if norm_val in seen:
            continue
        if cnt_3 < min_n:
            continue
        cnt_n  = vc_n.get(val, 0)
        rate_3 = cnt_3 / total_3
        rate_n = cnt_n / total_n if total_n > 0 else 0.0
        if cnt_n == 0 or (rate_3 - rate_n) >= diff_thr:
            extras.append(norm_val)
            seen.add(norm_val)
    return extras


def _add_tie_notes(doc, tie_notes):
    """ヒートマップのタイ注釈（※N）をヒートマップ直下に追記"""
    if not tie_notes:
        return
    for nid, panel, col, row, main, val_n, n, tied in tie_notes:
        all_vals = [(main, val_n)] + [(v, c) for v, c in tied]
        vals_str = ','.join(f'{v} ({c}/{n})' for v, c in all_vals)
        p = doc.add_paragraph()
        run = p.add_run(f'※{nid}【{panel}】{col} / {row}：{vals_str}')
        run.font.size = Pt(10)
        _set_jpfont(run, 'Meiryo UI')
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _add_opponent_line(doc, df_section):
    """セクションの対象大学を赤字で1行追加"""
    opps = sorted(df_section['OPPONENT'].unique())
    p = doc.add_paragraph()
    run_label = p.add_run('対象大学：')
    run_label.bold = True
    run_val = p.add_run('、'.join(opps))
    run_val.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _build_dist_zones(yd_zones):
    """UI ヤード設定 → アナライザ用 DIST_ZONES リストを生成"""
    circles = '①②③④⑤⑥'
    result = []
    for i, (name, d_min, d_max) in enumerate(yd_zones):
        c = circles[i] if i < len(circles) else f'({i+1})'
        range_str = f'{d_min}〜' if d_max >= 99 else f'{d_min}〜{d_max}'
        result.append((f'{c} {name} ({range_str})', d_min, d_max))
    return result


def _build_normal_field_cols(yd_zones):
    """Normal 用フィールドヒートマップ列定義を生成"""
    cols = [('1st\n10', 1, 1, 999)]
    for name, d_min, d_max in yd_zones:
        range_str = f'{d_min}~' if d_max >= 99 else f'{d_min}~{d_max}'
        cols.append((f'2nd\n{range_str}', 2, d_min, d_max))
    return cols


def _build_third_field_cols(yd_zones):
    """3rd 用フィールドヒートマップ列定義を生成"""
    cols = []
    for name, d_min, d_max in yd_zones:
        range_str = f'{d_min}~' if d_max >= 99 else f'{d_min}~{d_max}'
        cols.append((f'3rd/4th\n{range_str}', (3, 4), d_min, d_max))
    return cols


# ── レポート生成 ─────────────────────────────────────────────
def generate_word_report(df, situation_opps=None, yd_zones=None):
    # PERSONNEL_NORM が古い df に存在しない場合はここで付与
    if 'PERSONNEL_NORM' not in df.columns:
        import re as _re
        df = df.copy()
        df['PERSONNEL_NORM'] = df['PERSONNEL'].apply(
            lambda v: _re.sub(r'\.0$', '', str(v).strip())
            if str(v).strip().lower() not in ('nan', '') else ''
        )
    doc = Document()

    # ヒートマップ挿入幅：ページ全幅（マージン込み）でぎりぎりまで大きく
    _sec = doc.sections[0]
    _HM_W = _sec.page_width

    # ブックマークIDカウンターをリセット（複数回生成時の重複防止）
    _BM_ID[0] = 0

    # ドキュメント全体の Normal スタイルで段落後スペース・フォントを設定
    _normal_style = doc.styles['Normal']
    _normal_style.paragraph_format.space_after  = Pt(0)
    _normal_style.paragraph_format.space_before = Pt(0)
    # 日本語フォントをドキュメントレベルで 游ゴシック Medium に設定
    _normal_style.font.name = '游ゴシック Medium'
    _nrPr = _normal_style.element.get_or_add_rPr()
    _nFonts = _nrPr.get_or_add_rFonts()
    _nFonts.set(qn('w:eastAsia'), '游ゴシック Medium')

    # ── 表紙 ──
    doc.add_heading('スカウティングレポート', 0)
    doc.add_paragraph(f'生成日：{date.today().strftime("%Y年%m月%d日")}')
    doc.add_paragraph(f'総プレー数（全大学）：{len(df)}')
    _toc_runs = _add_toc(doc)
    _add_page_number_footer(doc)

    _all_opps = sorted(df['OPPONENT'].unique())

    # ヤードゾーン定義（UI設定 or デフォルト）
    _DEFAULT_YD = [
        ('SHORT',    1,   3),
        ('MIDDLE ①', 4,   6),
        ('MIDDLE ②', 7,  10),
        ('LONG',    11,  99),
    ]
    _yd = yd_zones if yd_zones is not None else _DEFAULT_YD
    _dist_zones       = _build_dist_zones(_yd)
    _normal_field_cols = _build_normal_field_cols(_yd)
    _third_field_cols  = _build_third_field_cols(_yd)

    # DIST_ZONE 列を df に追加（auto_analyzer で使用）
    def _assign_dist_zone(dist_val):
        for zone_label, d_min, d_max in _dist_zones:
            if d_min <= dist_val <= d_max:
                return zone_label
        return ''
    df = df.copy()
    df['DIST_ZONE'] = pd.to_numeric(df['DIST'], errors='coerce').apply(
        lambda d: _assign_dist_zone(int(d)) if pd.notna(d) else ''
    )

    # ── シチュエーション df を事前構築（並列処理の準備）────────────────
    _sel_n = situation_opps.get('normal', _all_opps) if situation_opps else _all_opps
    _sel_3 = situation_opps.get('3rd',    _all_opps) if situation_opps else _all_opps
    df_n = az.filter_normal(df[df['OPPONENT'].isin(_sel_n)])
    df_3 = az.filter_3rd(df[df['OPPONENT'].isin(_sel_3)])

    # ゾーン別結果・分析テーブル（pandas 操作なので高速）
    zone_results  = az.analyze_3rd_zones(df_3, dist_zones=_dist_zones)
    _front_tbl_n  = az.analyze_front(df_n)
    _sign_tbl_n   = az.analyze_sign_d(df_n)
    _sign_tbl_n5  = az.analyze_sign_d(df_n, min_count=5)   # 5プレー以上のみ（特徴量用）
    _sign_tbl_3_all = az.analyze_sign_d(df_3)              # 3rd全体SIGN表
    _cov_n, _comp3_n = az.analyze_coverage(df_n)

    # PERSONNEL フィルタ済み df（ヒートマップ & プレー数表示で共用）
    _pers_dfs_n = {g: az.filter_personnel(df_n, v) for g, v in az.PERSONNEL_GROUPS}
    _pers_dfs_3 = {g: az.filter_personnel(df_3, v) for g, v in az.PERSONNEL_GROUPS}

    # ── RedZone / 2MIN df をプール投入前に構築（並列化のため前出し）──────────
    _sel_r = situation_opps.get('red',  _all_opps) if situation_opps else _all_opps
    _sel_2 = situation_opps.get('2min', _all_opps) if situation_opps else _all_opps
    df_r = az.filter_redzone(df[df['OPPONENT'].isin(_sel_r)])
    df_2 = az.filter_2min(df[df['OPPONENT'].isin(_sel_2)])
    _pers_dfs_r = {g: az.filter_personnel(df_r, v) for g, v in az.PERSONNEL_GROUPS}

    # RedZone / 2MIN 分析テーブルを事前計算（純 pandas・高速）
    _front_tbl_r     = az.analyze_front(df_r)
    _cov_r, _comp3_r = az.analyze_coverage(df_r)
    _manzono_r       = az.analyze_manzono(df_r)
    _pkg_tbl_r       = az.analyze_redzone_packages(df_r)
    _pkg_map_r       = az.build_redzone_pkg_map(df_r)
    _pkg_map_rn      = {_norm(k): v for k, v in _pkg_map_r.items()}
    _sign_tbl_2      = az.analyze_sign_d(df_2)
    _cov_2, _comp3_2 = az.analyze_coverage(df_2)

    # RedZone ゾーン別テーブルも事前計算
    _rz_zone_dfs      = {}
    _pkg_tbl_rz_dict  = {}
    _pkg_map_rzn_dict = {}
    for _rz_lbl_pre, _rz_lo_pre, _rz_hi_pre in az.REDZONE_YARD_LN_ZONES:
        _df_rz_pre = df_r[df_r['YARD LN'].between(_rz_lo_pre, _rz_hi_pre)]
        _rz_zone_dfs[_rz_lbl_pre]      = _df_rz_pre
        _pkg_tbl_rz_dict[_rz_lbl_pre]  = az.analyze_redzone_packages(_df_rz_pre)
        _m_rz = az.build_redzone_pkg_map(_df_rz_pre)
        _pkg_map_rzn_dict[_rz_lbl_pre] = {_norm(k): v for k, v in _m_rz.items()}

    # ── 自動考察を並列実行（auto_analyzer のみ。matplotlib はメインスレッド専用）──
    # _add_tokuchou は _norm(raw_val) を consider_fn に渡すため、
    # 辞書キーも _norm 済みの値を使う（例: 'N' → 'read'）
    def _safe_result(fut):
        """Future の結果を安全に取得。例外は [] を返す"""
        try:
            return fut.result()
        except Exception:
            return []

    def _submit_items(pool, fn, tbl, df_sub):
        """テーブル第1列の各値（_norm済み）をキーにして future dict を作る"""
        return {_norm(str(v)): pool.submit(fn, v, df_sub)
                for v in tbl.iloc[:, 0]
                if pd.notna(v) and str(v).strip()}

    _pool = _TPE(max_workers=8)

    # Normal 自動考察ヘッダー
    _aa_fh_n = _pool.submit(aa.consider_front_header,     df_n)
    _aa_sh_n = _pool.submit(aa.consider_sign_header,      df_n)
    _aa_ch_n = _pool.submit(aa.consider_coverage_header,  df_n)
    _aa_3v_n = _pool.submit(aa.consider_3var_offform_cov, df_n)
    _aa_n3c  = _pool.submit(aa.consider_n3_comparison, df_n, df_3)
    # Normal SIGNシチュエーション分析（5プレー以上のSIGN全件、Normal専用ブロック）
    _aa_sign_sit_nh = _pool.submit(aa.consider_sign_situation_header, df_n, 20, 5, 0.60)
    # 3rd SIGNシチュエーション分析（全3rdプレー横断）
    _aa_sign_sit_3h = _pool.submit(aa.consider_sign_situation_header, df_3)
    # 3rd SIGN 割合ヘッダー（Normalとの違いを含む）
    _aa_sh_3g = _pool.submit(aa.consider_sign_header, df_3, df_n)

    # Normal 自動考察アイテム（_norm済みキー）
    _aa_fi_n = _submit_items(_pool, aa.consider_front_item, _front_tbl_n, df_n)
    _aa_si_n = _submit_items(_pool, aa.consider_sign_item,  _sign_tbl_n,  df_n)
    _aa_ci_n = _submit_items(_pool, aa.consider_cov_item,   _cov_n,       df_n)

    # 3rd ゾーン 自動考察
    _aa_zfh: dict = {}; _aa_zsh: dict = {}; _aa_zch: dict = {}
    _aa_zfi: dict = {}; _aa_zsi: dict = {}; _aa_zci: dict = {}
    for _zn, _zdata in zone_results.items():
        _df_z = _zdata['df']
        if len(_df_z) == 0:
            continue
        _aa_zfh[_zn] = _pool.submit(aa.consider_front_header,    _df_z)
        _aa_zsh[_zn] = _pool.submit(aa.consider_sign_header,     _df_z)
        _aa_zch[_zn] = _pool.submit(aa.consider_coverage_header, _df_z)
        _aa_zfi[_zn] = _submit_items(_pool, aa.consider_front_item, _zdata['front'],    _df_z)
        _aa_zsi[_zn] = _submit_items(_pool, aa.consider_sign_item,  _zdata['sign'],     _df_z)
        _aa_zci[_zn] = _submit_items(_pool, aa.consider_cov_item,   _zdata['coverage'], _df_z)

    # ── RedZone / 2MIN 自動考察をプールに投入（シャットダウン前）─────────────
    _aa_rz_fh = _pool.submit(aa.consider_redzone_front_header,
                             df_r, df_normal=df_n, df_3rd=df_3)
    _aa_rz_ch = _pool.submit(aa.consider_redzone_cov_header,
                             df_r, df_normal=df_n, df_3rd=df_3)
    _aa_rz_ph = _pool.submit(aa.consider_redzone_packages_header, df_r)
    _aa_rz_pi = {
        _norm(str(v)): _pool.submit(
            aa.consider_redzone_package_item,
            _norm(str(v)), _pkg_map_rn.get(_norm(str(v))), df_r)
        for v in (_pkg_tbl_r.iloc[:, 0] if len(_pkg_tbl_r) else [])
        if pd.notna(v) and str(v).strip()
    }
    _aa_rz_zone_ph: dict = {}
    _aa_rz_zone_dd: dict = {}
    _aa_rz_zone_pi: dict = {}
    for _rz_lbl2, _df_rz2, _tbl_rz2, _mapn2 in (
        (_lbl, _rz_zone_dfs[_lbl], _pkg_tbl_rz_dict[_lbl], _pkg_map_rzn_dict[_lbl])
        for _lbl in _rz_zone_dfs
    ):
        _aa_rz_zone_ph[_rz_lbl2] = _pool.submit(
            aa.consider_redzone_packages_header, _df_rz2, df_overall=df_r)
        _aa_rz_zone_dd[_rz_lbl2] = _pool.submit(
            aa.consider_redzone_packages_downdist, _df_rz2)
        _aa_rz_zone_pi[_rz_lbl2] = {
            _norm(str(v)): _pool.submit(
                aa.consider_redzone_package_item,
                _norm(str(v)), _mapn2.get(_norm(str(v))), _df_rz2)
            for v in (_tbl_rz2.iloc[:, 0] if len(_tbl_rz2) else [])
            if pd.notna(v) and str(v).strip()
        }
    _aa_2m_sh = _pool.submit(aa.consider_2min_sign_header, df_2)
    _aa_2m_ch = _pool.submit(aa.consider_2min_cov_header,  df_2)

    # ── フィールドグリッド計算をプールで先行実行（純 pandas・スレッドセーフ）──
    _GRID_N   = _pool.submit(az.analyze_field_grid, df_n,
                             situation='normal', field_cols=_normal_field_cols)
    _GRID_3   = _pool.submit(az.analyze_field_grid, df_3,
                             situation='third',  field_cols=_third_field_cols)
    _GRID_RN  = _pool.submit(az.analyze_field_grid, df_r,
                             situation='normal', yard_zones=az.REDZONE_YARD_LN_ZONES)
    _GRID_R3  = _pool.submit(az.analyze_field_grid, df_r,
                             situation='third',  yard_zones=az.REDZONE_YARD_LN_ZONES)
    _GRID_PN  = {g: _pool.submit(az.analyze_field_grid, _pers_dfs_n[g],
                                  situation='normal', field_cols=_normal_field_cols)
                 for g, _ in az.PERSONNEL_GROUPS}
    _GRID_P3  = {g: _pool.submit(az.analyze_field_grid, _pers_dfs_3[g],
                                  situation='third',  field_cols=_third_field_cols)
                 for g, _ in az.PERSONNEL_GROUPS}
    _GRID_PRN = {g: _pool.submit(az.analyze_field_grid, _pers_dfs_r[g],
                                  situation='normal', yard_zones=az.REDZONE_YARD_LN_ZONES)
                 for g, _ in az.PERSONNEL_GROUPS}
    _GRID_PR3 = {g: _pool.submit(az.analyze_field_grid, _pers_dfs_r[g],
                                  situation='third',  yard_zones=az.REDZONE_YARD_LN_ZONES)
                 for g, _ in az.PERSONNEL_GROUPS}

    # 投入完了。メインスレッドはドキュメント構築を続けながらバックグラウンドで演算が走る
    _pool.shutdown(wait=False)

    _SIT_LABELS = [
        ('normal', 'Normal Situation'),
        ('3rd',    '3rd Situation'),
        ('red',    'Red Zone'),
        ('2min',   '2MIN'),
    ]

    doc.add_paragraph()
    p_note = doc.add_paragraph()
    run_note = p_note.add_run('【大学選択】各シチュエーションの対象大学・除外大学')
    run_note.bold = True
    run_note.font.size = Pt(12)

    for sit_key, sit_label in _SIT_LABELS:
        sel = situation_opps.get(sit_key, _all_opps) if situation_opps else _all_opps
        exc = [o for o in _all_opps if o not in sel]

        p = doc.add_paragraph()
        p.add_run(f'{sit_label}　').bold = True

        run_sel = p.add_run(f'選択：{"、".join(sel)}')
        run_sel.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        run_sel.bold = True

        p.add_run('　')
        run_exc = p.add_run(f'除外：{"、".join(exc)}' if exc else '除外：なし')
        run_exc.bold = True

    # ════════════════════════════════════════
    # 1. Normal Situation
    # ════════════════════════════════════════
    _add_h1(doc, '1. Normal Situation', bookmark='bm_normal')
    _p = doc.add_paragraph()
    _r = _p.add_run('フィルタ条件：DN=1,2 ／ 2MIN除外 ／ RedZone（YARD LN 1〜25）除外')
    _r.bold = True; _r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.add_paragraph(f'該当プレー数：{len(df_n)}')
    _add_opponent_line(doc, df_n)

    # 1-1 Run Defense
    _add_h2(doc, '1-1. Run Defense')

    _add_table(doc, _front_tbl_n, 'DEF FRONT 割合',
               auto_bullets=_safe_result(_aa_fh_n))
    _add_tokuchou(doc, '各Frontの特徴', _front_tbl_n, named_count=5, total_count=6,
                  consider_fn=lambda v: _safe_result(_aa_fi_n[v]) if v in _aa_fi_n else None)

    doc.add_page_break()
    _add_table(doc, _sign_tbl_n, 'SIGN 割合',
               auto_bullets=_safe_result(_aa_sh_n))
    # SIGN 要素状況分析（専用ブロック：5プレー以上のSIGN全件）
    _sit_nh = _safe_result(_aa_sign_sit_nh)
    if _sit_nh:
        _add_auto_section(doc, _sit_nh)
    _add_tokuchou(doc, '各Signの特徴', _sign_tbl_n5,
                  named_count=len(_sign_tbl_n5), total_count=len(_sign_tbl_n5),
                  consider_fn=lambda v: _safe_result(_aa_si_n[v]) if v in _aa_si_n else None)

    doc.add_page_break()
    _p = doc.add_paragraph()
    _r = _p.add_run('OFF FORM ごとのSIGN割合')
    _r.bold = True
    _r.font.size = Pt(18)
    _add_off_form_sign_d_paired(doc, az.analyze_off_form_sign_d(df_n))

    # 1-2 Pass Defense
    doc.add_page_break()
    _add_h2(doc, '1-2. Pass Defense')

    _add_coverage_with_comp3(doc, _cov_n, _comp3_n, 'パスカバー割合（COVERAGE）',
                             auto_bullets=_safe_result(_aa_ch_n))
    _add_tokuchou(doc, '各Coverの特徴', _cov_n, named_count=5, total_count=6,
                  consider_fn=lambda v: _safe_result(_aa_ci_n[v]) if v in _aa_ci_n else None)

    doc.add_page_break()
    _p = doc.add_paragraph()
    _r = _p.add_run('OFF FORM ごとのCOVERAGE割合')
    _r.bold = True
    _r.font.size = Pt(18)
    _add_off_form_paired(doc, az.analyze_off_form_coverage(df_n), three_var_dict=_safe_result(_aa_3v_n))

    # 1-3. フィールド型ヒートマップ（matplotlib はメインスレッドで生成）
    doc.add_page_break()
    _add_h2(doc, '1-3. フィールド型ヒートマップ')
    _add_h3(doc, 'Normal')
    _hm_buf, _hm_ties = cg.generate_field_heatmap(
        _GRID_N.result(), situation_label='Normal Situation')
    doc.add_picture(_hm_buf, width=_HM_W)
    _set_picture_float_square(doc)
    _add_tie_notes(doc, _hm_ties)

    # 1-4. PERSONNEL別ヒートマップ
    doc.add_page_break()
    _add_h2(doc, '1-4. PERSONNEL別ヒートマップ')
    p_pers_n = doc.add_paragraph('パーソネルごとのヒートマップ')
    p_pers_n.runs[0].font.size = Pt(12)
    _set_jpfont(p_pers_n.runs[0], 'Meiryo UI')
    p_pers_n.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    for _pn_idx, (grp_label, grp_vals) in enumerate(az.PERSONNEL_GROUPS):
        if _pn_idx > 0:
            doc.add_page_break()
        _add_h3(doc, f'{grp_label}per')
        doc.add_paragraph(f'該当プレー数：{len(_pers_dfs_n[grp_label])}')
        _hm_buf_pn, _hm_ties_pn = cg.generate_field_heatmap(
            _GRID_PN[grp_label].result(),
            situation_label=f'PERSONNEL {grp_label} — Normal',
            )
        doc.add_picture(_hm_buf_pn, width=_HM_W)
        _set_picture_float_square(doc)
        _add_tie_notes(doc, _hm_ties_pn)
        doc.add_paragraph()

    # ════════════════════════════════════════
    # 2. 3rd Situation
    # ════════════════════════════════════════
    _add_h1(doc, '2. 3rd Situation', bookmark='bm_3rd')
    _p = doc.add_paragraph()
    _r = _p.add_run('フィルタ条件：DN=3,4 ／ 2MIN除外 ／ RedZone除外')
    _r.bold = True; _r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    doc.add_paragraph(f'該当プレー数：{len(df_3)}')
    _add_opponent_line(doc, df_3)

    # Normal vs 3rd 横断比較（ヤード非依存）を冒頭に表示
    _add_auto_section(doc, _safe_result(_aa_n3c))

    # 3rd全体 SIGN割合表（Normalとの違いを上に、SIGNシチュエーション分析を下に）
    _add_table(doc, _sign_tbl_3_all, '3rd SIGN 割合',
               auto_bullets=_safe_result(_aa_sh_3g))
    _sit_3h = _safe_result(_aa_sign_sit_3h)
    if _sit_3h:
        _add_auto_section(doc, _sit_3h)

    # ヤードゾーンごとの詳細分析
    for _zone_idx, (zone_name, data) in enumerate(zone_results.items()):
        n = data['n']
        _zn_parts = zone_name.split(' ', 1)
        _zn_display = f'{_zn_parts[0]} 3rd & {_zn_parts[1]}' if len(_zn_parts) == 2 else zone_name
        if _zone_idx > 0:
            doc.add_page_break()
        _add_h2(doc, f'{_zn_display}  （{n}プレー）')
        if n == 0:
            doc.add_paragraph('  該当プレーなし')
            continue

        # 事前計算済みのフューチャーを取得
        _zfi = _aa_zfi.get(zone_name, {})
        _zsi = _aa_zsi.get(zone_name, {})
        _zci = _aa_zci.get(zone_name, {})

        # ゾーン名サフィックス（例：「3rd & SHORT (1〜3)」）
        _zone_suffix = f'（{"3rd & " + _zn_parts[1] if len(_zn_parts) == 2 else zone_name}）'

        # FRONT分析
        _add_table(doc, data['front'], f'DEF FRONT 割合{_zone_suffix}',
                   auto_bullets=_safe_result(_aa_zfh[zone_name]) if zone_name in _aa_zfh else None)
        _front_top = [_norm(str(r.iloc[0])) for _, r in data['front'].head(5).iterrows()
                      if not pd.isna(r.iloc[0])]
        _extra_front = _3rd_notable_names(data['df'], df_n, 'DEF_FRONT', _front_top)
        _add_tokuchou(doc, '各Frontの特徴', data['front'], named_count=5,
                      consider_fn=lambda v, _fi=_zfi: _safe_result(_fi[v]) if v in _fi else None,
                      show_fig_placeholder=False, extra_names=_extra_front)

        # SIGN分析
        doc.add_page_break()
        _add_table(doc, data['sign'], f'SIGN 割合{_zone_suffix}',
                   auto_bullets=_safe_result(_aa_zsh[zone_name]) if zone_name in _aa_zsh else None)
        _sign_top = [_norm(str(r.iloc[0])) for _, r in data['sign'].head(5).iterrows()
                     if not pd.isna(r.iloc[0])]
        _extra_sign = _3rd_notable_names(data['df'], df_n, 'SIGN_D', _sign_top)
        _add_tokuchou(doc, '各Signの特徴', data['sign'], named_count=5,
                      consider_fn=lambda v, _si=_zsi: _safe_result(_si[v]) if v in _si else None,
                      show_fig_placeholder=False, extra_names=_extra_sign)

        # COVERAGE分析
        doc.add_page_break()
        _add_coverage_with_comp3(doc, data['coverage'], data['comp3'], f'COVERAGE 割合{_zone_suffix}',
                                 auto_bullets=_safe_result(_aa_zch[zone_name]) if zone_name in _aa_zch else None)
        _cov_top = [_norm(str(r.iloc[0])) for _, r in data['coverage'].head(5).iterrows()
                    if not pd.isna(r.iloc[0])]
        _extra_cov = _3rd_notable_names(data['df'], df_n, 'COVERAGE_NORM', _cov_top)
        _add_tokuchou(doc, '各Coverの特徴', data['coverage'], named_count=5,
                      consider_fn=lambda v, _ci=_zci: _safe_result(_ci[v]) if v in _ci else None,
                      show_fig_placeholder=False, extra_names=_extra_cov)

        # パッケージ
        pkg = data['packages']
        if len(pkg) > 0:
            _add_table(doc, pkg, 'よく出るパッケージ（3プレー以上）', header_color='FCE4D6')
        else:
            doc.add_paragraph('  よく出るパッケージ：なし')

    # 2-3. フィールド型ヒートマップ
    doc.add_page_break()
    _add_h2(doc, '2-3. フィールド型ヒートマップ')
    _add_h3(doc, 'フィールド型ヒートマップ')
    _hm_buf_3, _hm_ties_3 = cg.generate_field_heatmap(
        _GRID_3.result(), situation_label='3rd Situation — 3rd')
    doc.add_picture(_hm_buf_3, width=_HM_W)
    _set_picture_float_square(doc)
    _add_tie_notes(doc, _hm_ties_3)
    doc.add_paragraph()

    # 2-4. PERSONNEL別ヒートマップ
    doc.add_page_break()
    _add_h2(doc, '2-4. PERSONNEL別ヒートマップ')
    p_pers_3 = doc.add_paragraph('パーソネルごとのヒートマップ')
    p_pers_3.runs[0].font.size = Pt(12)
    _set_jpfont(p_pers_3.runs[0], 'Meiryo UI')
    p_pers_3.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    for _p3_idx, (grp_label, grp_vals) in enumerate(az.PERSONNEL_GROUPS):
        if _p3_idx > 0:
            doc.add_page_break()
        _add_h3(doc, f'{grp_label}per')
        doc.add_paragraph(f'該当プレー数：{len(_pers_dfs_3[grp_label])}')
        _hm_buf_p3, _hm_ties_p3 = cg.generate_field_heatmap(
            _GRID_P3[grp_label].result(),
            situation_label=f'PERSONNEL {grp_label} — 3rd',
            )
        doc.add_picture(_hm_buf_p3, width=_HM_W)
        _set_picture_float_square(doc)
        _add_tie_notes(doc, _hm_ties_p3)
        doc.add_paragraph()

    # ════════════════════════════════════════
    # 3. Red Zone
    # ════════════════════════════════════════
    _add_h1(doc, '3. Red Zone', bookmark='bm_red')
    _p = doc.add_paragraph()
    _r = _p.add_run('フィルタ条件：YARD LN 1〜25 ／ 2MIN除外')
    _r.bold = True; _r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph(f'該当プレー数：{len(df_r)}')
    _add_opponent_line(doc, df_r)

    # Run Defense
    _add_h2(doc, '3-1. Run Defense')
    _add_table(doc, _front_tbl_r, 'DEF FRONT 割合',
               auto_bullets=_safe_result(_aa_rz_fh))

    # Pass Defense
    doc.add_page_break()
    _add_h2(doc, '3-2. Pass Defense')
    _add_coverage_with_comp3(doc, _cov_r, _comp3_r, 'COVERAGE 割合',
                             auto_bullets=_safe_result(_aa_rz_ch))
    # Man系 / Zone系 比率
    _add_table(doc, _manzono_r, 'Man / Zone 比率')

    # パッケージ割合（全体 + ヤードゾーン別）
    doc.add_page_break()
    _add_h2(doc, '3-3. パッケージ割合')
    _add_table(doc, _pkg_tbl_r, '全体',
               auto_bullets=_safe_result(_aa_rz_ph))
    _add_tokuchou(doc, '各パッケージの特徴', _pkg_tbl_r,
                  named_count=min(5, len(_pkg_tbl_r)),
                  total_count=min(5, len(_pkg_tbl_r)),
                  consider_fn=lambda v, _pi=_aa_rz_pi: _safe_result(_pi[v]) if v in _pi else None,
                  show_fig_placeholder=True)
    for _rz_label, _rz_min, _rz_max in az.REDZONE_YARD_LN_ZONES:
        _df_rz      = _rz_zone_dfs[_rz_label]
        _tbl_rz     = _pkg_tbl_rz_dict[_rz_label]
        _mapn_rz    = _pkg_map_rzn_dict[_rz_label]
        _zone_pi_rz = _aa_rz_zone_pi[_rz_label]
        _rz_dd_bullets = _safe_result(_aa_rz_zone_dd[_rz_label])
        if _rz_dd_bullets:
            _add_auto_section(doc, _rz_dd_bullets)
        _add_table(doc, _tbl_rz,
                   f'{_rz_label}（{len(_df_rz)}プレー）',
                   auto_bullets=_safe_result(_aa_rz_zone_ph[_rz_label]))
        _add_tokuchou(doc, '各パッケージの特徴', _tbl_rz,
                      named_count=min(5, len(_tbl_rz)),
                      total_count=min(5, len(_tbl_rz)),
                      consider_fn=lambda v, _pi=_zone_pi_rz: _safe_result(_pi[v]) if v in _pi else None,
                      show_fig_placeholder=True)

    # 3-4. フィールド型ヒートマップ（Normal: 1st/2nd）
    doc.add_page_break()
    _add_h2(doc, '3-4. フィールド型ヒートマップ')
    _add_h3(doc, 'Normal（1st / 2nd）Redzone')
    _hm_buf_rn, _hm_ties_rn = cg.generate_field_heatmap(
        _GRID_RN.result(),
        situation_label='Normal（1st / 2nd）Redzone',
        redzone_mode=True,
    )
    doc.add_picture(_hm_buf_rn, width=_HM_W)
    _set_picture_float_square(doc)
    _add_tie_notes(doc, _hm_ties_rn)
    doc.add_paragraph()

    # 3rd
    doc.add_page_break()
    _add_h3(doc, '3rd Redzone')
    _hm_buf_r3, _hm_ties_r3 = cg.generate_field_heatmap(
        _GRID_R3.result(),
        situation_label='3rd Redzone',
        redzone_mode=True,
    )
    doc.add_picture(_hm_buf_r3, width=_HM_W)
    _set_picture_float_square(doc)
    _add_tie_notes(doc, _hm_ties_r3)
    doc.add_paragraph()

    # 3-5. PERSONNEL別ヒートマップ
    doc.add_page_break()
    _add_h2(doc, '3-5. PERSONNEL別ヒートマップ')
    p_pers_r = doc.add_paragraph('パーソネルごとのヒートマップ')
    p_pers_r.runs[0].font.size = Pt(12)
    _set_jpfont(p_pers_r.runs[0], 'Meiryo UI')
    p_pers_r.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    for _pr_idx, (grp_label, grp_vals) in enumerate(az.PERSONNEL_GROUPS):
        if _pr_idx > 0:
            doc.add_page_break()
        _add_h3(doc, f'{grp_label}per  Normal（1st / 2nd）Redzone')
        doc.add_paragraph(f'該当プレー数：{len(_pers_dfs_r[grp_label])}')
        _hm_buf_prn, _hm_ties_prn = cg.generate_field_heatmap(
            _GRID_PRN[grp_label].result(),
            situation_label=f'PERSONNEL {grp_label} Normal（1st / 2nd）Redzone',
            redzone_mode=True,
        )
        doc.add_picture(_hm_buf_prn, width=_HM_W)
        _set_picture_float_square(doc)
        _add_tie_notes(doc, _hm_ties_prn)
        doc.add_paragraph()

        doc.add_page_break()
        _add_h3(doc, f'{grp_label}per  3rd Redzone')
        _hm_buf_pr3, _hm_ties_pr3 = cg.generate_field_heatmap(
            _GRID_PR3[grp_label].result(),
            situation_label=f'PERSONNEL {grp_label} 3rd Redzone',
            redzone_mode=True,
        )
        doc.add_picture(_hm_buf_pr3, width=_HM_W)
        _set_picture_float_square(doc)
        _add_tie_notes(doc, _hm_ties_pr3)
        doc.add_paragraph()

    # ════════════════════════════════════════
    # 4. 2MIN
    # ════════════════════════════════════════
    _add_h1(doc, '4. 2MIN', bookmark='bm_2min')
    _p = doc.add_paragraph()
    _r = _p.add_run('フィルタ条件：2MIN = Y（DN・DIST・YARD LN によらず）')
    _r.bold = True; _r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph(f'該当プレー数：{len(df_2)}')
    _add_opponent_line(doc, df_2)

    _add_table(doc, _sign_tbl_2, 'SIGN 割合',
               auto_bullets=_safe_result(_aa_2m_sh))

    doc.add_page_break()
    _add_coverage_with_comp3(doc, _cov_2, _comp3_2, 'COVERAGE 割合',
                             auto_bullets=_safe_result(_aa_2m_ch))

    # ════════════════════════════════════════
    # 5. 選手の特徴
    # ════════════════════════════════════════
    _add_h1(doc, '5. 選手の特徴', bookmark='bm_players')

    # ── 目次ページ番号をPythonで確定 ──
    _bm_names = [bm for _, bm in _TOC_ENTRIES]
    _page_nums = _estimate_page_nums(doc, _bm_names)
    for bm_name, run in _toc_runs.items():
        n = _page_nums.get(bm_name)
        run.text = f'{n}ページ' if n is not None else '-ページ'

    # ── BytesIO に保存して返す ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
